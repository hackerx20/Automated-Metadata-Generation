import os
import re
import logging
from typing import Dict, Optional, Tuple, List
import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import io
import fitz  # PyMuPDF for advanced PDF processing
from datetime import datetime
import hashlib

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Document processing module for extracting text from various file formats
    """
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from document based on file type
        
        Args:
            file_path: Path to the document file
            file_type: MIME type of the file
            
        Returns:
            Extracted text content
        """
        try:
            if 'pdf' in file_type.lower():
                return self.extract_text_from_pdf(file_path)
            elif 'word' in file_type.lower() or 'docx' in file_type.lower():
                return self.extract_text_from_docx(file_path)
            elif 'text' in file_type.lower() or 'txt' in file_type.lower():
                return self.extract_text_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Enhanced PDF text extraction with comprehensive layout handling
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text content with improved structure preservation
        """
        text = ""
        
        try:
            # Use PyMuPDF for advanced PDF processing
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text with layout preservation
                page_text = self._extract_structured_text(page)
                
                # Handle multi-column layouts
                if self._is_multi_column(page):
                    page_text = self._process_multi_column_text(page)
                
                # Extract headers and footers separately
                headers_footers = self._extract_headers_footers(page)
                
                text += f"\n--- Page {page_num + 1} ---\n"
                if headers_footers['header']:
                    text += f"HEADER: {headers_footers['header']}\n"
                
                text += page_text + "\n"
                
                if headers_footers['footer']:
                    text += f"FOOTER: {headers_footers['footer']}\n"
                
                # Extract tables if present
                tables = self._extract_tables_from_page(page)
                if tables:
                    text += "\nTABLES:\n" + "\n".join(tables) + "\n"
            
            doc.close()
            
            # If extraction yielded very little text, try OCR
            if len(text.strip()) < 100:
                ocr_text = self._ocr_pdf(file_path)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            # Fallback to pdfplumber
            try:
                text = self._fallback_pdf_extraction(file_path)
            except Exception as fallback_error:
                logger.error(f"Fallback extraction failed: {fallback_error}")
                # Final fallback to OCR
                try:
                    text = self._ocr_pdf(file_path)
                except Exception as ocr_error:
                    text = f"Error extracting text: {str(e)}"
        
        return self.preprocess_text(text)
    
    def _extract_structured_text(self, page) -> str:
        """Extract text with structure preservation using PyMuPDF"""
        try:
            # Get text with layout information
            text_dict = page.get_text("dict")
            structured_text = ""
            
            for block in text_dict["blocks"]:
                if "lines" in block:  # Text block
                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            line_text += span["text"]
                        structured_text += line_text + "\n"
                        
            return structured_text
        except Exception:
            # Fallback to simple text extraction
            return page.get_text()
    
    def _is_multi_column(self, page) -> bool:
        """Detect if page has multi-column layout"""
        try:
            text_dict = page.get_text("dict")
            x_positions = []
            
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        if line["spans"]:
                            x_positions.append(line["spans"][0]["bbox"][0])
            
            # If we have significantly different x-positions, likely multi-column
            if len(set(int(x/50)*50 for x in x_positions)) > 2:
                return True
            return False
        except Exception:
            return False
    
    def _process_multi_column_text(self, page) -> str:
        """Process multi-column text with proper ordering"""
        try:
            text_dict = page.get_text("dict")
            columns = {}
            
            for block in text_dict["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        if line["spans"]:
                            x_pos = int(line["spans"][0]["bbox"][0] / 100) * 100
                            if x_pos not in columns:
                                columns[x_pos] = []
                            
                            line_text = "".join(span["text"] for span in line["spans"])
                            y_pos = line["spans"][0]["bbox"][1]
                            columns[x_pos].append((y_pos, line_text))
            
            # Sort columns by x-position and text by y-position
            result = ""
            for x_pos in sorted(columns.keys()):
                column_text = sorted(columns[x_pos], key=lambda x: x[0])
                result += "\n".join(text for _, text in column_text) + "\n\n"
            
            return result
        except Exception:
            return page.get_text()
    
    def _extract_headers_footers(self, page) -> Dict[str, str]:
        """Extract headers and footers from page"""
        try:
            rect = page.rect
            header_rect = fitz.Rect(0, 0, rect.width, rect.height * 0.1)
            footer_rect = fitz.Rect(0, rect.height * 0.9, rect.width, rect.height)
            
            header_text = page.get_textbox(header_rect).strip()
            footer_text = page.get_textbox(footer_rect).strip()
            
            return {
                'header': header_text if len(header_text) > 3 else "",
                'footer': footer_text if len(footer_text) > 3 else ""
            }
        except Exception:
            return {'header': "", 'footer': ""}
    
    def _extract_tables_from_page(self, page) -> List[str]:
        """Extract tables from page using PyMuPDF"""
        tables = []
        try:
            # Find tables using PyMuPDF
            table_finder = page.find_tables()
            for table in table_finder:
                table_data = table.extract()
                if table_data:
                    # Convert table to string representation
                    table_str = ""
                    for row in table_data:
                        if row:  # Skip empty rows
                            table_str += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                    if table_str.strip():
                        tables.append(table_str)
        except Exception:
            pass
        
        return tables
    
    def _fallback_pdf_extraction(self, file_path: str) -> str:
        """Fallback PDF extraction using pdfplumber"""
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.error(f"Fallback extraction error: {e}")
        return text
    
    def extract_pdf_metadata(self, file_path: str) -> Dict:
        """Extract comprehensive PDF metadata"""
        metadata = {
            'pdf_version': None,
            'page_count': 0,
            'creation_date': None,
            'modification_date': None,
            'author': None,
            'creator': None,
            'producer': None,
            'title': None,
            'subject': None,
            'keywords': None,
            'security_settings': {},
            'bookmarks': [],
            'embedded_files': [],
            'form_fields': [],
            'annotations': [],
            'page_sizes': [],
            'is_encrypted': False,
            'has_signatures': False
        }
        
        try:
            # Use PyMuPDF for comprehensive metadata extraction
            doc = fitz.open(file_path)
            
            # Basic metadata
            pdf_metadata = doc.metadata
            metadata.update({
                'author': pdf_metadata.get('author', ''),
                'creator': pdf_metadata.get('creator', ''),
                'producer': pdf_metadata.get('producer', ''),
                'title': pdf_metadata.get('title', ''),
                'subject': pdf_metadata.get('subject', ''),
                'keywords': pdf_metadata.get('keywords', ''),
                'creation_date': pdf_metadata.get('creationDate', ''),
                'modification_date': pdf_metadata.get('modDate', ''),
                'page_count': len(doc),
                'is_encrypted': doc.needs_pass
            })
            
            # Extract bookmarks/outline
            try:
                toc = doc.get_toc()
                metadata['bookmarks'] = [{'level': level, 'title': title, 'page': page} 
                                       for level, title, page in toc]
            except Exception:
                pass
            
            # Extract page sizes
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                rect = page.rect
                metadata['page_sizes'].append({
                    'page': page_num + 1,
                    'width': rect.width,
                    'height': rect.height
                })
            
            # Check for embedded files
            try:
                embedded_files = doc.embfile_names()
                metadata['embedded_files'] = embedded_files or []
            except Exception:
                pass
            
            # Check for form fields
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                try:
                    widgets = page.widgets()
                    for widget in widgets:
                        metadata['form_fields'].append({
                            'page': page_num + 1,
                            'type': widget.field_type_string,
                            'name': widget.field_name
                        })
                except Exception:
                    pass
            
            # Check for annotations
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                try:
                    annotations = page.annots()
                    for annot in annotations:
                        annot_dict = annot.info
                        metadata['annotations'].append({
                            'page': page_num + 1,
                            'type': annot_dict.get('type', ''),
                            'content': annot_dict.get('content', '')
                        })
                except Exception:
                    pass
            
            doc.close()
            
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['page_count'] = len(pdf_reader.pages)
                    metadata['is_encrypted'] = pdf_reader.is_encrypted
                    
                    if pdf_reader.metadata:
                        pdf_info = pdf_reader.metadata
                        metadata.update({
                            'author': pdf_info.get('/Author', ''),
                            'creator': pdf_info.get('/Creator', ''),
                            'producer': pdf_info.get('/Producer', ''),
                            'title': pdf_info.get('/Title', ''),
                            'subject': pdf_info.get('/Subject', ''),
                        })
            except Exception as fallback_error:
                logger.error(f"Fallback metadata extraction failed: {fallback_error}")
        
        return metadata
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            return text
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
            raise
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file with encoding detection
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
            raise
    
    def _ocr_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using OCR (for scanned documents)
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            OCR extracted text
        """
        try:
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Convert page to image
                    page_image = page.to_image(resolution=300)
                    
                    # Convert PIL image to format pytesseract can use
                    img_bytes = io.BytesIO()
                    page_image.save(img_bytes, format='PNG')
                    img_bytes.seek(0)
                    
                    # Perform OCR
                    page_text = pytesseract.image_to_string(
                        Image.open(img_bytes),
                        config='--psm 3'  # Page segmentation mode
                    )
                    
                    if page_text.strip():
                        text += page_text + "\n"
            
            return text
            
        except Exception as e:
            logger.warning(f"OCR failed for {file_path}: {str(e)}")
            return ""
    
    def preprocess_text(self, text: str) -> str:
        """
        Clean and preprocess extracted text
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned and preprocessed text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace and newlines
        text = re.sub(r'\n+', '\n', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\|\~\`]', '', text)
        
        # Remove excessive punctuation
        text = re.sub(r'([.!?]){2,}', r'\1', text)
        
        # Clean up common OCR errors
        text = re.sub(r'\b[A-Z]{1}\s+(?=[A-Z])', '', text)  # Remove single letters
        text = re.sub(r'\s+([.!?,:;])', r'\1', text)  # Fix spacing before punctuation
        
        # Normalize quotes
        text = re.sub(r'[""''`]', '"', text)
        
        # Strip and return
        return text.strip()
    
    def detect_file_type(self, file_path: str) -> str:
        """
        Detect file type based on extension and content
        
        Args:
            file_path: Path to file
            
        Returns:
            Detected file type
        """
        _, ext = os.path.splitext(file_path.lower())
        
        type_mapping = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain'
        }
        
        return type_mapping.get(ext, 'unknown')
    
    def get_document_structure(self, text: str) -> Dict:
        """
        Identify document structure (headings, sections, etc.)
        
        Args:
            text: Document text
            
        Returns:
            Dictionary containing structure information
        """
        structure = {
            'headings': [],
            'sections': [],
            'paragraphs': [],
            'has_toc': False
        }
        
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Detect headings (lines that are short, capitalized, or numbered)
            if self._is_heading(line):
                structure['headings'].append({
                    'text': line,
                    'line_number': i,
                    'level': self._get_heading_level(line)
                })
            
            # Detect sections
            if len(line.split()) < 10 and any(word in line.lower() for word in 
                                           ['chapter', 'section', 'part', 'introduction', 'conclusion']):
                structure['sections'].append({
                    'text': line,
                    'line_number': i
                })
        
        # Split into paragraphs
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        structure['paragraphs'] = paragraphs[:10]  # First 10 paragraphs
        
        # Check for table of contents
        structure['has_toc'] = self._detect_toc(text)
        
        return structure
    
    def _is_heading(self, line: str) -> bool:
        """Check if a line is likely a heading"""
        if len(line) > 100:  # Too long to be a heading
            return False
        
        # Check for numbered headings
        if re.match(r'^\d+\.?\s+[A-Z]', line):
            return True
        
        # Check for all caps headings
        if line.isupper() and len(line.split()) <= 8:
            return True
        
        # Check for title case headings
        words = line.split()
        if len(words) <= 8 and sum(1 for w in words if w[0].isupper() if w) > len(words) * 0.7:
            return True
        
        return False
    
    def _get_heading_level(self, line: str) -> int:
        """Determine heading level (1-6)"""
        # Check for numbered headings
        match = re.match(r'^(\d+)\.', line)
        if match:
            return min(int(match.group(1)), 6)
        
        # Based on length and formatting
        if line.isupper():
            return 1
        elif len(line.split()) <= 3:
            return 2
        else:
            return 3
    
    def _detect_toc(self, text: str) -> bool:
        """Detect if document has a table of contents"""
        toc_indicators = [
            'table of contents', 'contents', 'index',
            r'\d+\s+\.\.\.\s+\d+',  # Page numbers with dots
            r'chapter \d+.*\d+$'     # Chapter with page number
        ]
        
        for indicator in toc_indicators:
            if re.search(indicator, text.lower(), re.MULTILINE):
                return True
        
        return False
